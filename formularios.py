from docx import Document
from docx.enum.section import WD_ORIENT

# Crear documento Word
doc = Document()

# Establecer orientación horizontal (paisaje)
section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# Título principal
doc.add_heading("Miembros nuevos", level=1)

# Agregar los puntos explicativos del 1 al 5
items = [
    "1) ¿Tienen hnos/as ministrantes amorosos/as que les estén ministrando?",
    "2) ¿Tienen recomendación para el templo para efectuar bautismos?",
    "3) ¿Tienen una asignación o llamamiento?",
    "4) ¿La maestra/o de miembros nuevos los/as conoce y los/as tiene asistiendo a su clase?",
    "5) ¿Tiene una recomendación para recibir su bendición patriarcal?"
]

doc.add_paragraph("Puntos para verificar:", style='List Number')
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph("")  # Espacio antes de la tabla

# Encabezados de la tabla
headers = ['Hermanos:'] + items + ['Observaciones']

# Crear la tabla con encabezados
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells

for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Crear 30 filas para nombres con opciones Sí/No
for _ in range(30):
    row_cells = table.add_row().cells
    row_cells[0].text = ''  # Nombre del miembro
    for i in range(1, len(headers) - 1):
        row_cells[i].text = 'Sí ☐  No ☐'
    row_cells[-1].text = ''  # Observaciones

# Guardar el documento
file_path = "Miembros nuevos.docx"
doc.save(file_path)

file_path
